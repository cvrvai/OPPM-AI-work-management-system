"""
File parser — extracts text from uploaded binary files for LLM consumption.

Supported formats:
- .xlsx / .xls  → openpyxl (sheets → markdown tables)
- .pdf          → pdfplumber (page-by-page text extraction)
- .docx         → python-docx (paragraphs + tables)
- .csv          → stdlib csv (markdown table)
- text fallback → decode as UTF-8
"""

import csv
import io
import logging
from dataclasses import dataclass

logger = logging.getLogger(__name__)

MAX_EXTRACTED_CHARS = 30_000
MAX_FILE_BYTES = 10 * 1024 * 1024  # 10 MB


@dataclass
class ParseResult:
    """Result from parsing a file."""
    text: str
    truncated: bool = False
    error: str | None = None


def _truncate(text: str) -> tuple[str, bool]:
    """Truncate text to MAX_EXTRACTED_CHARS if needed."""
    if len(text) <= MAX_EXTRACTED_CHARS:
        return text, False
    return text[:MAX_EXTRACTED_CHARS] + "\n\n… (content truncated — file exceeds extraction limit)", True


def _parse_xlsx(content: bytes) -> ParseResult:
    """Extract text from Excel (.xlsx) files as markdown tables."""
    try:
        import openpyxl
    except ImportError:
        return ParseResult(text="", error="openpyxl not installed — cannot parse Excel files")

    try:
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    except Exception as e:
        logger.warning("Failed to open Excel file: %s", e)
        return ParseResult(text="", error=f"Cannot open Excel file: {e}")

    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        parts.append(f"### Sheet: {sheet_name}\n")

        # Build markdown table
        header = rows[0]
        col_names = [str(c) if c is not None else "" for c in header]
        parts.append("| " + " | ".join(col_names) + " |")
        parts.append("| " + " | ".join(["---"] * len(col_names)) + " |")

        for row in rows[1:]:
            cells = [str(c) if c is not None else "" for c in row]
            parts.append("| " + " | ".join(cells) + " |")

        parts.append("")  # blank line between sheets

    wb.close()

    if not parts:
        return ParseResult(text="(empty spreadsheet — no data found)")

    text = "\n".join(parts)
    text, truncated = _truncate(text)
    return ParseResult(text=text, truncated=truncated)


def _parse_pdf(content: bytes) -> ParseResult:
    """Extract text from PDF files page-by-page."""
    try:
        import pdfplumber
    except ImportError:
        return ParseResult(text="", error="pdfplumber not installed — cannot parse PDF files")

    try:
        pdf = pdfplumber.open(io.BytesIO(content))
    except Exception as e:
        logger.warning("Failed to open PDF file: %s", e)
        return ParseResult(text="", error=f"Cannot open PDF file: {e}")

    parts: list[str] = []
    for i, page in enumerate(pdf.pages, 1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = "(page extraction failed)"
        if page_text.strip():
            parts.append(f"--- Page {i} ---\n{page_text.strip()}")

    pdf.close()

    if not parts:
        return ParseResult(text="(PDF contains no extractable text — may be a scanned image)")

    text = "\n\n".join(parts)
    text, truncated = _truncate(text)
    return ParseResult(text=text, truncated=truncated)


def _parse_docx(content: bytes) -> ParseResult:
    """Extract text from Word (.docx) files — paragraphs and tables."""
    try:
        import docx
    except ImportError:
        return ParseResult(text="", error="python-docx not installed — cannot parse Word files")

    try:
        doc = docx.Document(io.BytesIO(content))
    except Exception as e:
        logger.warning("Failed to open Word file: %s", e)
        return ParseResult(text="", error=f"Cannot open Word file: {e}")

    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Extract tables as markdown
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            # Add header separator after first row
            col_count = len(table.rows[0].cells) if table.rows else 0
            if len(rows) > 0:
                header_sep = "| " + " | ".join(["---"] * col_count) + " |"
                rows.insert(1, header_sep)
            parts.append("\n".join(rows))

    if not parts:
        return ParseResult(text="(empty document — no text found)")

    text = "\n\n".join(parts)
    text, truncated = _truncate(text)
    return ParseResult(text=text, truncated=truncated)


def _parse_csv(content: bytes) -> ParseResult:
    """Extract text from CSV files as a markdown table."""
    try:
        text_content = content.decode("utf-8", errors="replace")
    except Exception as e:
        return ParseResult(text="", error=f"Cannot decode CSV file: {e}")

    try:
        reader = csv.reader(io.StringIO(text_content))
        rows = list(reader)
    except Exception as e:
        return ParseResult(text="", error=f"Cannot parse CSV file: {e}")

    if not rows:
        return ParseResult(text="(empty CSV file)")

    parts: list[str] = []
    header = rows[0]
    parts.append("| " + " | ".join(header) + " |")
    parts.append("| " + " | ".join(["---"] * len(header)) + " |")

    for row in rows[1:]:
        # Pad or trim to match header length
        cells = row + [""] * (len(header) - len(row))
        parts.append("| " + " | ".join(cells[:len(header)]) + " |")

    text = "\n".join(parts)
    text, truncated = _truncate(text)
    return ParseResult(text=text, truncated=truncated)


def _parse_text(content: bytes) -> ParseResult:
    """Fallback: decode as UTF-8 text."""
    try:
        text = content.decode("utf-8", errors="replace")
    except Exception as e:
        return ParseResult(text="", error=f"Cannot decode file: {e}")

    text, truncated = _truncate(text)
    return ParseResult(text=text, truncated=truncated)


# ── Dispatch table ──

_PARSERS: dict[str, type] = {}  # populated below

_EXT_MAP = {
    ".xlsx": _parse_xlsx,
    ".xls": _parse_xlsx,
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
    ".doc": _parse_docx,
    ".csv": _parse_csv,
}


def parse_file(filename: str, content: bytes) -> ParseResult:
    """Parse a file and return extracted text.

    Dispatches to the appropriate parser based on file extension.
    Falls back to UTF-8 text decoding for unknown file types.

    Args:
        filename: Original filename (used for extension detection).
        content: Raw file bytes.

    Returns:
        ParseResult with extracted text, truncation status, and any error.
    """
    if len(content) > MAX_FILE_BYTES:
        return ParseResult(
            text="",
            error=f"File too large ({len(content) / 1024 / 1024:.1f} MB). Maximum is {MAX_FILE_BYTES / 1024 / 1024:.0f} MB.",
        )

    if not content:
        return ParseResult(text="(empty file)")

    ext = ""
    dot_idx = filename.rfind(".")
    if dot_idx >= 0:
        ext = filename[dot_idx:].lower()

    parser = _EXT_MAP.get(ext, _parse_text)
    try:
        return parser(content)
    except Exception as e:
        logger.warning("File parser failed for '%s': %s", filename, e)
        return ParseResult(text="", error=f"Failed to parse file: {e}")
